# -*- coding: utf-8 -*-
"""
Date created: Tue Mar  7 00:08:08 2017
Author: S. Hutchins

Description: Use the docx module to create a Microsoft Word template.

"""

# Modules used
from docx import Document
from docx.shared import Pt
#------------------------------------------------------------------------------
# Open/Create a file
doc = Document()

# Create a run object.
run = doc.add_paragraph().add_run()
font = run.font

# Set font type and size
font.name = 'Calibri'
font.size = Pt(12)

# Add a heading. Levels 0 - 9. This is level 1 and the default.
doc.add_heading('Example', level=1)

# Add a paragraph to that file
paragraph = doc.add_paragraph('First paragraph.')

# This is a subheading
doc.add_heading('Subheading', level=2)

# Create a bulleted list
paragraph = doc.add_paragraph('First paragraph.', style = 'ListBullet')

# Add a page break
doc.add_page_break()

# Save the file
doc.save('test.docx')
#------------------------------------------------------------------------------

